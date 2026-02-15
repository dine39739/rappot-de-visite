import streamlit as st
from datetime import date
from fpdf import FPDF
from PIL import Image
import os
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH  # <--- Indispensable pour l'alignement

# --- CONFIGURATION DE LA PAGE ---
st.set_page_config(page_title="Tech-Report Pro", layout="wide", page_icon="🏗️")

# --- INITIALISATION DES VARIABLES (SESSION STATE) ---
if 'participants' not in st.session_state:
    st.session_state.participants = []
if 'sections' not in st.session_state:
    st.session_state.sections = [{'titre': '', 'description': '', 'photos': []}]

# --- STYLE CSS POUR LE RENDU ---
st.markdown("""
    <style>
    .main { background-color: #f5f7f9; }
    .stButton>button { width: 100%; border-radius: 5px; height: 3em; }
    .section-container { border: 1px solid #ddd; padding: 20px; border-radius: 10px; margin-bottom: 20px; background-color: white; }
    </style>
    """, unsafe_allow_html=True)

st.title("🏗️ Générateur de Rapport Technique")
st.info("Remplissez les sections ci-dessous. Vous pouvez ajouter autant de participants et de sections que nécessaire.")

# --- ÉTAPE 1 : INFOS GÉNÉRALES ---
with st.expander("📌 Informations du Chantier", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        client_name = st.text_input("Nom du Client / Projet", placeholder="ex: Résidence Les Palmiers")
        adresse = st.text_input("Adresse de l'intervention")
    with col2:
        date_visite = st.date_input("Date de la visite", date.today())
        technicien = st.text_input("Technicien responsable")

# --- ÉTAPE 2 : PARTICIPANTS ---
st.header("👥 Participants")
if st.button("➕ Ajouter un participant"):
    st.session_state.participants.append({"nom": "", "tel": "", "email": ""})

for i, p in enumerate(st.session_state.participants):
    with st.container():
        c1, c2, c3, c4 = st.columns([3, 2, 3, 1])
        p['nom'] = c1.text_input(f"Nom & Prénom", value=p['nom'], key=f"p_nom_{i}")
        p['tel'] = c2.text_input(f"Téléphone", value=p['tel'], key=f"p_tel_{i}")
        p['email'] = c3.text_input(f"Email", value=p['email'], key=f"p_email_{i}")
        if c4.button("🗑️", key=f"del_p_{i}"):
            st.session_state.participants.pop(i)
            st.rerun()

# --- ÉTAPE 3 : SECTIONS DU RAPPORT ---
st.header("📝 Corps du Rapport")

for idx, sec in enumerate(st.session_state.sections):
    with st.container():
        st.markdown(f"**Section {idx + 1}**")
        sec['titre'] = st.text_input("Titre de la section", value=sec['titre'], key=f"sec_titre_{idx}", placeholder="ex: Constatations en toiture")
        sec['description'] = st.text_area("Observations détaillées", value=sec['description'], key=f"sec_desc_{idx}")
        
        # Gestion des photos pour cette section
        sec['photos'] = st.file_uploader(f"Ajouter des photos (Section {idx+1})", 
                                         accept_multiple_files=True, 
                                         type=['png', 'jpg', 'jpeg'], 
                                         key=f"sec_img_{idx}")
        
        if len(st.session_state.sections) > 1:
            if st.button(f"❌ Supprimer la section {idx+1}", key=f"del_sec_{idx}"):
                st.session_state.sections.pop(idx)
                st.rerun()
        st.divider()

if st.button("➕ Ajouter une Section de travail"):
    st.session_state.sections.append({'titre': '', 'description': '', 'photos': []})

# --- FONCTION DE GÉNÉRATION PDF ---
def generate_pdf():
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # --- 1. CHARGEMENT DE LA POLICE UNICODE ---
    # Assurez-vous que le fichier .ttf est bien sur votre GitHub
    pdf.add_font("DejaVu", '', 'DejaVuSans.ttf')
    pdf.set_font("DejaVu", '', 12)

    # --- 2. EN-TÊTE AVEC LOGO ---
    # pdf.image(nom_du_fichier, x, y, largeur)
    # Si le fichier logo.png existe, on l'affiche
    if os.path.exists("logo.png"):
        pdf.image("logo.png", x=10, y=8, w=30)
    
    pdf.ln(20) # Saut de ligne après le logo

    # --- 3. TITRE ENCADRÉ (Bleu foncé, texte blanc) ---
    # Couleurs RGB : Bleu foncé (0, 51, 102), Blanc (255, 255, 255)
    pdf.set_fill_color(0, 51, 102)  # Couleur du fond de l'encadré
    pdf.set_text_color(255, 255, 255) # Couleur du texte
    pdf.set_font("DejaVu", '', 18)
    
    # Cell(largeur, hauteur, texte, bordure, retour ligne, alignement, remplissage)
    pdf.cell(0, 15, "RAPPORT D'INTERVENTION TECHNIQUE", ln=True, align='C', fill=True)
    
    # --- 4. RÉINITIALISATION POUR LE RESTE DU TEXTE ---
    pdf.set_text_color(0, 0, 0) # On repasse en noir
    pdf.set_font("DejaVu", '', 11)
    pdf.ln(5)
    
    # Infos générales (sous le titre)
    pdf.set_font("helvetica", '', 10)
    pdf.cell(0, 7, f"Client : {client_name}", ln=True)
    pdf.cell(0, 7, f"Adresse : {adresse}", ln=True)
    pdf.cell(0, 7, f"Date : {date_visite} | Technicien : {technicien}", ln=True)
    pdf.ln(10)

    # --- 5. SECTION PARTICIPANTS ---
    if st.session_state.participants:
        pdf.set_font("helvetica", '', 12)
        pdf.set_fill_color(230, 230, 230) # Gris très clair
        pdf.cell(0, 10, " PERSONNES PRÉSENTES", ln=True, fill=True)
        pdf.set_font("DejaVu", '', 10)
        for p in st.session_state.participants:
            pdf.cell(0, 8, f"• {p['nom']} (Tél: {p['tel']} | Email: {p['email']})", ln=True)
        pdf.ln(10)

    # --- 6. CORPS DU RAPPORT ---
    for sec in st.session_state.sections:
        if sec['titre']:
            # Titre de section stylisé (souligné bleu)
            pdf.set_font("helvetica", '', 14)
            pdf.set_text_color(0, 51, 102)
            pdf.cell(0, 10, sec['titre'].upper(), ln=True)
            pdf.set_draw_color(0, 51, 102)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
            pdf.ln(2)
            
            # Description
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("helvetica", '', 11)
            pdf.multi_cell(0, 7, sec['description'])
            pdf.ln(5)

            # Photos
            if sec['photos']:
                # On organise les photos par 2 par ligne pour gagner de la place
                col_width = 90
                for i, img_file in enumerate(sec['photos']):
                    try:
                        img = Image.open(img_file)
                        if img.mode in ("RGBA", "P"):
                            img = img.convert("RGB")
                        
                        temp_path = f"temp_{idx}_{i}_{img_file.name}"
                        img.save(temp_path)
                        
                        # Gestion de l'espace pour ne pas couper l'image en bas de page
                        if pdf.get_y() > 220:
                            pdf.add_page()
                        
                        pdf.image(temp_path, w=col_width)
                        pdf.ln(5)
                        os.remove(temp_path)
                    except Exception as e:
                        st.error(f"Erreur photo : {e}")
            pdf.ln(10)

    return pdf.output()

# --- BOUTON FINAL ---
st.divider()
if st.button("🚀 GÉNÉRER LE RAPPORT PDF"):
    if not client_name or not technicien:
        st.warning("Veuillez remplir au moins le nom du client et du technicien.")
    else:
        with st.spinner("Création du PDF en cours..."):
            pdf_data = generate_pdf()
            st.success("✅ Votre rapport est prêt !")
            st.download_button(
                label="⬇️ Télécharger le Rapport (PDF)",
                data=bytes(pdf_data),
                file_name=f"Rapport_{client_name}_{date_visite}.pdf",
                mime="application/pdf"
            )

# --- PROCHAINE ÉTAPE : GOOGLE DRIVE ---
# Note : Pour lier à Drive, il faudra configurer les "Secrets" dans Streamlit Cloud.

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# --- CONFIGURATION DRIVE ---
# Remplacez par l'ID de votre dossier Drive (il est dans l'URL de votre dossier)
FOLDER_ID = "1izwpTbS9x5fUI2a0UWQVWmlG3XcKNEDn" 

def upload_to_drive(pdf_bytes, filename):
    try:
        info = st.secrets["gcp_service_account"]
        creds = service_account.Credentials.from_service_account_info(info)
        service = build('drive', 'v3', credentials=creds)

        # Configuration du fichier
        file_metadata = {
            'name': filename,
            'parents': [FOLDER_ID]
        }
        
        fh = io.BytesIO(pdf_bytes)
        media = MediaIoBaseUpload(fh, mimetype='application/pdf')

        # L'ASTUCE : On force le fichier à ne pas utiliser le quota du robot
        # En partageant le dossier avec le robot en tant qu'éditeur, 
        # le fichier hérite de la propriété du dossier parent (le vôtre).
        file = service.files().create(
            body=file_metadata, 
            media_body=media, 
            fields='id',
            supportsAllDrives=True # Important pour la gestion des quotas
        ).execute()
        
        return file.get('id')
    except Exception as e:
        st.error(f"Erreur Drive : {e}")
        return None

import json

# Préparation des données
donnees_brouillon = {
    "client": client_name,
    "adresse": adresse,
    "technicien": technicien,
    "participants": st.session_state.participants,
    "sections": [{"titre": s["titre"], "description": s["description"]} for s in st.session_state.sections]
}

# Conversion en texte
json_string = json.dumps(donnees_brouillon, indent=4)

st.sidebar.header("💾 Persistance locale")
st.sidebar.download_button(
    label="📥 Sauvegarder l'état actuel",
    data=json_string,
    file_name=f"brouillon_{client_name}.json",
    mime="application/json",
    help="Télécharge un petit fichier qui contient tout votre texte actuel."
)
# --- DANS VOTRE BOUTON DE GÉNÉRATION FINAL ---

if st.button("🚀 GÉNÉRER ET ENVOYER LE RAPPORT"):
    if not client_name:
        st.error("Veuillez saisir le nom du client.")
    else:
        with st.spinner("Génération du PDF et synchronisation Drive..."):
            pdf_data = generate_pdf()
            pdf_bytes = bytes(pdf_data)
            
            # 1. Sauvegarde sur Drive
            filename = f"Rapport_{client_name}_{date_visite}.pdf"
            file_id = upload_to_drive(pdf_bytes, filename)
            
            if file_id: "1izwpTbS9x5fUI2a0UWQVWmlG3XcKNEDn"
            st.success(f"✅ Rapport sauvegardé sur Google Drive !")
            
            # 2. Proposer quand même le téléchargement local
            st.download_button(
                label="⬇️ Télécharger une copie locale",
                data=pdf_bytes,
                file_name=filename,
                mime="application/pdf"
            )

import urllib.parse

# Préparation du lien mailto
sujet = f"Rapport d'intervention : {client_name}"
corps = f"Bonjour,\n\nVeuillez trouver ci-joint le rapport pour l'intervention du {date_visite}.\n\nCordialement,"
# Encodage pour les espaces et caractères spéciaux
mail_link = f"mailto:?subject={urllib.parse.quote(sujet)}&body={urllib.parse.quote(corps)}"

st.markdown(f'<a href="{mail_link}" target="_blank"><button style="width:100%; height:3em; background-color:#0078d4; color:white; border:none; border-radius:5px;">📧 Ouvrir dans Outlook</button></a>', unsafe_allow_html=True)


# --- FONCTION GÉNÉRATION WORD ---
def generate_word():
    doc = Document()
    
    # 1. Récupération sécurisée des variables
    nom_client = st.session_state.get('client_name', 'Client Inconnu').upper()
    nom_tech = st.session_state.get('technicien', 'Non renseigné')
    visite_date = str(st.session_state.get('date_visite', ''))
    lieu = st.session_state.get('adresse', 'Non renseignée')
    
    # 2. Titre Principal
    title = doc.add_heading(f"RAPPORT : {nom_client}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. En-tête Infos
    p = doc.add_paragraph()
    p.add_run("Date de la visite : ").bold = True
    p.add_run(f"{visite_date}\n")
    p.add_run("Technicien : ").bold = True
    p.add_run(f"{nom_tech}\n")
    p.add_run("Adresse : ").bold = True
    p.add_run(f"{lieu}")

    # 4. Participants
    doc.add_heading("Participants", level=1)
    parts = st.session_state.get('participants', [])
    if isinstance(parts, list):
        for part in parts:
            if isinstance(part, dict):
                nom_p = part.get('nom', '')
                soc_p = part.get('societe', '')
                doc.add_paragraph(f"• {nom_p} ({soc_p})", style='List Bullet')

    # 5. Sections et Photos
    doc.add_heading("Constats et Photos", level=1)
    sections = st.session_state.get('sections', [])
    
    for s in sections:
        doc.add_heading(s.get('titre', 'Sans titre'), level=2)
        doc.add_paragraph(s.get('description', ''))
        
        if s.get('image') is not None:
            try:
                image_bytes = s['image'].getvalue()
                image_stream = io.BytesIO(image_bytes)
                doc.add_picture(image_stream, width=Inches(4.0))
                doc.add_paragraph() 
            except Exception as e:
                p_err = doc.add_paragraph()
                p_err.add_run(f"[Image non insérée : {e}]").italic = True

    # --- ATTENTION : CES LIGNES DOIVENT ÊTRE DÉCALÉES DE 4 ESPACES ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer 

# --- SECTION EXPORT FINAL (INTERFACE) ---
# Ici on revient tout à gauche car on sort de la fonction
st.divider()
st.subheader("🏁 Finaliser le Rapport")

col_pdf, col_word = st.columns(2)



with col_word:
    if st.button("📝 Préparer le fichier Word"):
        word_buffer = generate_word() 
        if word_buffer:
            word_data = word_buffer.getvalue() 
            st.download_button(
                label="⬇️ Cliquer pour télécharger (.docx)",
                data=word_data,
                file_name=f"Rapport_{st.session_state.get('client_name', 'Export')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
# --- BARRE LATÉRALE : SAUVEGARDE ET RESTAURATION LOCALE ---

st.sidebar.header("💾 Gestion du Dossier")

save_data = {
    "client_name": st.session_state.cli_val,
    "adresse": st.session_state.adr_val,
    "technicien": st.session_state.tec_val,
    "date_visite": str(st.session_state.date_val),
    "participants": st.session_state.participants,
    "sections": images_to_base64(st.session_state.sections)
}
st.sidebar.download_button("📥 Télécharger JSON", json.dumps(save_data, indent=4), "sauvegarde.json")

uploaded = st.sidebar.file_uploader("📂 Charger un fichier JSON", type=["json"])

if uploaded and st.sidebar.button("♻️ RESTAURER LES DONNÉES"):
    data = json.load(uploaded)
    
    # 1. CETTE BOUCLE EST LA CORRECTION : 
    # Elle supprime les anciennes valeurs des cases (widgets)
    # 't_' pour les titres, 'd_' pour les descriptions
    for key in list(st.session_state.keys()):
        if key.startswith(('t_', 'd_', 'cli_val', 'adr_val', 'tec_val')):
            del st.session_state[key]
            
    # 2. Injection des données du fichier JSON dans la mémoire
    st.session_state.cli_val = data.get("client_name", "")
    st.session_state.adr_val = data.get("adresse", "")
    st.session_state.tec_val = data.get("technicien", "")
    
    # Restauration des sections (Toiture, etc.)
    st.session_state.sections = base64_to_images(data.get("sections", []))
    
    # 3. On force le redémarrage pour que Streamlit crée les cases vides
    # puis les remplisse avec les données ci-dessus  
    st.rerun()

# --- INTERFACE PRINCIPALE ---
st.title("🏗️ Tech-Report Pro")

with st.expander("📌 Informations du Chantier", expanded=True):
    col1, col2 = st.columns(2)
    st.text_input("Nom du Client / Projet", key="cli_val")
    st.text_input("Adresse de l'intervention", key="adr_val")
    st.text_input("Technicien responsable", key="tec_val")
    st.date_input("Date de la visite", key="date_val")

st.header("👥 Participants")
if st.button("➕ Ajouter un participant"):
    st.session_state.participants.append({"nom": "", "tel": "", "email": ""})
    st.rerun()

for i, p in enumerate(st.session_state.participants):
    c1, c2, c3, c4 = st.columns([3, 2, 3, 1])
    p['nom'] = c1.text_input(f"Nom {i}", value=p.get('nom',''), key=f"pnom_{i}")
    p['tel'] = c2.text_input(f"Tél {i}", value=p.get('tel',''), key=f"ptel_{i}")
    p['email'] = c3.text_input(f"Email {i}", value=p.get('email',''), key=f"pmail_{i}")
    if c4.button("🗑️", key=f"pdel_{i}"):
        st.session_state.participants.pop(i)
        st.rerun()

st.header("📝 Corps du Rapport")

# BOUCLE DES SECTIONS
for idx, sec in enumerate(st.session_state.sections):
    with st.container():
        st.subheader(f"Section {idx+1}")
        
        # On utilise value=sec.get(...) ET on met à jour la liste en direct
        # L'utilisation de value assure que même si la clé change, le texte reste
        st.session_state.sections[idx]['titre'] = st.text_input(
            f"Titre Section {idx+1}", 
            value=sec.get('titre', ''), 
            key=f"t_{idx}"
        )
        
        st.session_state.sections[idx]['description'] = st.text_area(
            f"Observations Section {idx+1}", 
            value=sec.get('description', ''), 
            key=f"d_{idx}",
            height=200
        )
        
        if sec.get('photos'):
            st.success(f"📸 {len(sec['photos'])} photo(s) chargée(s) pour cette section.")
        
        new_imgs = st.file_uploader(f"Ajouter/Remplacer photos S{idx+1}", accept_multiple_files=True, key=f"img_{idx}")
        if new_imgs:
            st.session_state.sections[idx]['photos'] = new_imgs
            
        if st.button(f"🗑️ Supprimer Section {idx+1}", key=f"sdel_{idx}"):
            st.session_state.sections.pop(idx)
            st.rerun()
        st.divider()

if st.button("➕ Ajouter une Section"):
    st.session_state.sections.append({'titre': '', 'description': '', 'photos': []})
    st.rerun()

# EXPORT
if st.button("📄 Générer le Rapport PDF"):
    pdf_res = generate_pdf()
    st.download_button("⬇️ Télécharger PDF", bytes(pdf_res) if not isinstance(pdf_res, str) else pdf_res.encode('latin-1'), f"Rapport_{st.session_state.cli_val}.pdf", "application/pdf")
